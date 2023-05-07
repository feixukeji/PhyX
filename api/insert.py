from lxml import etree
import latex2mathml.converter
from docx.shared import Pt

def latex_to_word(latex_code):
    # Latex代码转Word对象

    mathml = latex2mathml.converter.convert(latex_code.replace(r"^{\circ}", "°"))
    tree = etree.fromstring(mathml)
    xslt = etree.parse('api/MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def insert_data(docu, name, data, option):
    # 平均值、标准差、不确定度计算公式插入
    
    # 参数：
    # docu：文档对象
    # name：该物理量的名称和符号
    # data：该物理量的相关数据（包含平均值、标准差、不确定度的元组）
    # option：插入选项，"word"表示插入 Word 公式，"latex"表示插入 Latex 公式

    if option == "word":
        docu.add_paragraph(name + "的平均值")
        docu.add_paragraph()._element.append(latex_to_word(data.averagex2))
        docu.add_paragraph(name + "的标准差")
        docu.add_paragraph()._element.append(latex_to_word(data.sigmax2))
        docu.add_paragraph(name + "的B类不确定度")
        docu.add_paragraph()._element.append(latex_to_word(data.delta_bx2))
        docu.add_paragraph(name + "的展伸不确定度")
        docu.add_paragraph()._element.append(latex_to_word(data.uncx2))
        docu.add_paragraph()
    elif option == "latex":
        docu.add_paragraph(name + "的平均值")
        docu.add_paragraph(data.averagex)
        docu.add_paragraph(name + "的标准差")
        docu.add_paragraph(data.sigmax)
        docu.add_paragraph(name + "的B类不确定度")
        docu.add_paragraph(data.delta_bx)
        docu.add_paragraph(name + "的展伸不确定度")
        docu.add_paragraph(data.uncx)
        docu.add_paragraph()


def insert_data_lsm(docu, data, option):
    # 最小二乘法结果插入
    
    # 参数：
    # docu：文档对象
    # data：相关数据（包含斜率、截距、线性拟合的相关系数、斜率标准差、截距标准差的元组）
    # option：插入选项，"word"表示插入 Word 公式，"latex"表示插入 Latex 公式
    
    if option == "word":
        docu.add_paragraph("斜率")
        docu.add_paragraph()._element.append(latex_to_word(data.mx2))
        docu.add_paragraph("截距")
        docu.add_paragraph()._element.append(latex_to_word(data.bx2))
        docu.add_paragraph("线性拟合的相关系数")
        docu.add_paragraph()._element.append(latex_to_word(data.rx2))
        docu.add_paragraph("斜率的延伸不确定度")
        docu.add_paragraph()._element.append(latex_to_word(data.u_mx2))
        docu.add_paragraph("截距的延伸不确定度")
        docu.add_paragraph()._element.append(latex_to_word(data.u_bx2))
        font = docu.add_paragraph().add_run("这里只考虑斜率和截距的 A 类不确定度。由 x 物理量和 y 物理量的 B 类不确定度如何推出斜率的 B 类不确定度，这个问题大物教学中心并未回复。").font
        font.italic = True
        font.size = Pt(6.5)
        docu.add_paragraph()
    elif option == "latex":
        docu.add_paragraph("斜率")
        docu.add_paragraph(data.mx)
        docu.add_paragraph("截距")
        docu.add_paragraph(data.bx)
        docu.add_paragraph("线性拟合的相关系数")
        docu.add_paragraph(data.rx)
        docu.add_paragraph("斜率的延伸不确定度")
        docu.add_paragraph(data.u_mx)
        docu.add_paragraph("截距的延伸不确定度")
        docu.add_paragraph(data.u_bx)
        docu.add_paragraph()


def insert_data_com(docu, name, data, option):
    # 表达式及合成不确定度计算公式插入
    
    # 参数：
    # docu：文档对象
    # name：该物理量的名称和符号
    # data：该物理量的相关数据（包含表达式计算、合成不确定度的元组）
    # option：插入选项，"word"表示插入 Word 公式，"latex"表示插入 Latex 公式

    if option == "word":
        docu.add_paragraph(name)
        docu.add_paragraph()._element.append(latex_to_word(data.ansx2))
        docu.add_paragraph(name + "的延伸不确定度")
        docu.add_paragraph()._element.append(latex_to_word(data.uncx2))
        docu.add_paragraph(name + "的最终结果")
        docu.add_paragraph()._element.append(latex_to_word(data.finalx2))
        docu.add_paragraph()
    elif option == "latex":
        docu.add_paragraph(name)
        docu.add_paragraph(data.ansx)
        docu.add_paragraph(name + "的延伸不确定度")
        docu.add_paragraph(data.uncx)
        docu.add_paragraph(name + "的最终结果")
        docu.add_paragraph(data.finalx)
        docu.add_paragraph()